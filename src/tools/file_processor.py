"""File processing utilities for PDF, DOCX documents"""

import logging
from pathlib import Path
from typing import Optional, Tuple

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        from PyPDF2 import PdfReader

        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num}: {str(e)}")

        result = "\n".join(text)
        logger.info(f"Extracted {len(result)} characters from PDF")
        return result
    except ImportError:
        logger.error("PyPDF2 not installed")
        raise ValueError("PDF support requires: pip install PyPDF2")
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
        raise ValueError(f"Failed to extract PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document

        doc = Document(file_path)
        text = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)

        result = "\n".join(text)
        logger.info(f"Extracted {len(result)} characters from DOCX")
        return result
    except ImportError:
        logger.error("python-docx not installed")
        raise ValueError("DOCX support requires: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        raise ValueError(f"Failed to extract DOCX: {str(e)}")


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract text from any supported document file"""
    try:
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = file_path.suffix.lower()

        if file_type == "pdf" or file_extension == ".pdf":
            return extract_text_from_pdf(str(file_path))
        elif file_type == "docx" or file_extension in [".docx", ".doc"]:
            return extract_text_from_docx(str(file_path))
        elif file_type == "txt" or file_extension == ".txt":
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    except Exception as e:
        logger.error(f"File extraction failed: {str(e)}")
        raise


def validate_file_size(file_size_bytes: int, max_size_mb: int = 10) -> bool:
    """Validate file size"""
    max_bytes = max_size_mb * 1024 * 1024
    return file_size_bytes <= max_bytes


def get_file_info(file_path: str) -> dict:
    """Get file information"""
    try:
        file_path = Path(file_path)
        return {
            "name": file_path.name,
            "type": file_path.suffix.lower(),
            "size_bytes": file_path.stat().st_size,
            "size_mb": file_path.stat().st_size / (1024 * 1024),
        }
    except Exception as e:
        logger.error(f"Failed to get file info: {str(e)}")
        return {}
